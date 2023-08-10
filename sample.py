from docx.shared import Inches, Pt
from docx import Document
import datetime
import comtypes
import os

def gen_report(self, frame=None, num=None, save=None, result=None):
    """
    Generate PV report for storage purpose

    Data:
        model       : Model number
        jabil_sn    : Jabil serial number
        test_date   : test date
        failure     : issues
        time_x      : time used for PV

    Example:
    self.job_dict[report] = {
                            'model': 'xxxxx-xxxxx',
                            'jabil_sn': 'MYxxxxxxxx',
                            'test_date': '15th June 2023',
                            'failure_obj': 'YES/NO',
                            'failure_led': 'YES/NO',
                            'failure_scr': 'YES/NO',
                            'time_taken': '30 mins'
                            }
    """
    if frame is not None:
        image_name = 'test.jpg'
    elif save and result:
        filename = (r'C:/Users/yungng07/Documents/pdf-generator/results.txt')
        with open(filename, 'r') as file:
            lines = file.readlines()
        with open(filename, 'w') as file:
            for line in lines:
                if line.startswith(save):
                    line = save + " = " + str(result) + "\n"
                file.write(line)
    else:
        ['report']['test_date'] = str(datetime.datetime.now())
        doc = Document('template.docx')
        # Loop through paragraphs and replace placeholders with data from dictionary
        for paragraph in doc.paragraphs:
            # Reconstruct paragraph content
            new_paragraph_text = paragraph.text
            for key, value in ['report'].items():
                placeholder = '{{' + key + '}}'
                if placeholder in new_paragraph_text:
                    new_paragraph_text = new_paragraph_text.replace(placeholder, value)
                    print(new_paragraph_text)
            # Clear existing runs
            for run in paragraph.runs:
                run.text = ''
            # Add the new text with formatting
            run = paragraph.add_run(new_paragraph_text)
            font = run.font
            font.name = 'Times New Roman'
            font.size = Pt(12)
        # Loop through tables and replace placeholders in table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    new_cell_text = cell.text
                    for key, value in self.job_dict['report'].items():
                        placeholder = '{{' + key + '}}'
                        if placeholder in new_cell_text:
                            new_cell_text = new_cell_text.replace(placeholder, value)
                            print(new_cell_text)
                    # Clear existing cell text
                    cell.text = ""
                    # Add new text with formatting
                    run = cell.paragraphs[0].add_run(new_cell_text)
                    font = run.font
                    font.name = 'Times New Roman'
                    font.size = Pt(12)
        i = 0
        while i < len(doc.paragraphs):
            paragraph = doc.paragraphs[i]
            for placeholder, name in self.job_dict['report_images'].items():
                if placeholder in paragraph.text:
                    # Create a new paragraph
                    new_paragraph = doc.add_paragraph()
                    # Construct image path and insert image
                    image_path = os.path.join(r'C:/Users/yungng07/Documents/pdf-generator/image/', name)
                    new_paragraph.add_run().add_picture(image_path, width=Inches(6))

                    # Insert the new paragraph after the current one
                    p = paragraph._element.getparent()
                    p.insert(p.index(paragraph._element) + 1, new_paragraph._element)
                    # Delete the old paragraph
                    p.remove(paragraph._element)
                    break
            i += 1
        # Save as docx
        filename = ['report']['model'] + '_' + ['report']['jabil_sn'] + '.docx'
        doc.save(r'C:/Users/yungng07/Documents/pdf-generator/report/' + filename)

        # Step 4: Convert Word to PDF
        word = comtypes.client.CreateObject('Word.Application')
        filename_pdf = ['report']['model'] + '_' + ['report']['jabil_sn'] + '.pdf'
        doc_com = word.Documents.Open(r'C:/Users/yungng07/Documents/pdf-generator/report/' + filename)
        # PDF format: 17
        doc_com.SaveAs(r'C:/Users/yungng07/Documents/pdf-generator/report/' + filename_pdf, FileFormat=17)
        doc_com.Close()
        word.Quit()
